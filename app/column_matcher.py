"""
Column matching and data extraction utilities using pandas
"""
import pandas as pd
from pathlib import Path
from typing import Dict, List, Tuple, Any
import logging
import asyncio
from io import StringIO
from openai import OpenAI
import docx  # python-docx for Word files
import openpyxl  # For Excel files

from .models import ColumnMatchResult, ExtractedData
from .config import settings
from .utils import load_sme_csv_to_dataframe
from .prompts import CSV_CLEANING_PROMPT, DOCUMENT_EXTRACTION_PROMPT

logger = logging.getLogger(__name__)


# Static column definitions for each template
TEMPLATE_COLUMNS = {
    "ADX_ESG": [
        "Section / القسم",
        "Field (EN)",
        "الحقل (AR)",
        "Prev Year",
        "Current",
        "Target",
        "Unit",
        "Notes",
        "Applicability",
        "Input Type",
        "Options"
    ],
    "DIFC_ESG": [
        "Section / القسم",
        "Field (EN)",
        "الحقل (AR)",
        "Prev Year",
        "Current",
        "Target",
        "Unit",
        "Notes",
        "Applicability",
        "Input Type",
        "Options",
        "Evidence Required?",
        "Confidential?"
    ],
    "MOCCAE": [
        "Section (EN)",
        "Section (AR)",
        "Field (EN)",
        "الحقل (AR)",
        "Response / الإدخال",
        "Unit / الوحدة",
        "Prev Year / العام السابق",
        "Current / العام الحالي",
        "Target / الهدف",
        "Notes / ملاحظات",
        "Applicability",
        "Evidence Required?",
        "Confidential?"
    ],
    "SCHOOLS": [
        "Section / القسم",
        "Field (EN)",
        "الحقل (AR)",
        "Prev Year",
        "Current",
        "Target",
        "Unit",
        "Notes",
        "Applicability",
        "Input Type",
        "Options",
        "Evidence Required?",
        "Confidential?"
    ],
    "SME": [
        "Section / القسم",
        "Field (EN)",
        "الحقل (AR)",
        "Prev Year",
        "Current",
        "Target",
        "Unit",
        "Notes",
        "Applicability",
        "Input Type",
        "Options",
        "Evidence Required?",
        "Confidential?"
    ]
}

# Template file mapping
TEMPLATE_FILES = {
    "ADX_ESG": "ADX_ESG_Template_v2_10.csv",
    "DIFC_ESG": "DIFC_ESG_Template_v2_10.csv",
    "MOCCAE": "MOCCAE_Compliance_Template_v2_10.csv",
    "SCHOOLS": "Schools_Lite_Template_v2_10.csv",
    "SME": "SME_Lite_Template_v2_10.csv"
}


class ColumnMatcher:
    """Handles column matching between uploaded files and templates using Grok AI"""
    
    def __init__(self, template_name: str):
        """
        Initialize the column matcher with a template
        
        Args:
            template_name: Name of the template (e.g., 'ADX_ESG', 'DIFC_ESG', 'MOCCAE', 'SCHOOLS', 'SME')
        """
        self.template_name = template_name
        self.template_columns = TEMPLATE_COLUMNS.get(template_name)
        
        if self.template_columns is None:
            raise ValueError(f"Unknown template: {template_name}. Available templates: {list(TEMPLATE_COLUMNS.keys())}")
        
        # Get template file path
        template_file = TEMPLATE_FILES.get(template_name)
        if template_file:
            self.template_path = settings.TEMPLATES_DIR / template_file
        else:
            self.template_path = None
        
        # Initialize Grok API client
        self.grok_client = OpenAI(
            api_key=settings.GROK_API_KEY,
            base_url=settings.GROK_API_BASE,
        )
        
        logger.info(f"Initialized column matcher for template: {self.template_name}")
    
    def get_template_columns(self) -> List[str]:
        """Get the list of column names from the template"""
        return self.template_columns
    
    def load_template_sample(self) -> str:
        """
        Load sample rows from the template file
        
        Returns:
            String representation of first 5 rows from template
        """
        if not self.template_path or not self.template_path.exists():
            logger.warning(f"Template file not found: {self.template_path}")
            return "Template sample not available"
        
        try:
            df = pd.read_csv(self.template_path, encoding='utf-8')
            # Get first 5 rows
            sample_df = df.head(5)
            # Convert to CSV string
            return sample_df.to_csv(index=False)
        except Exception as e:
            logger.error(f"Error loading template sample: {e}")
            return "Template sample not available"
    
    def load_uploaded_file(self, file_path: Path) -> pd.DataFrame:
        """
        Load uploaded file (CSV or Excel) into a DataFrame
        
        Args:
            file_path: Path to the uploaded file
            
        Returns:
            DataFrame containing the uploaded data
        """
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.csv':
                # Use robust CSV loader for uploaded CSV files
                df = load_sme_csv_to_dataframe(
                    str(file_path),
                    preserve_brackets=True,
                    merge_excess_into_notes=True,
                    encoding='utf-8',
                    verbose=False
                )
            elif file_extension in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            logger.info(f"Loaded uploaded file: {file_path.name}")
            return df
        except Exception as e:
            logger.error(f"Error loading uploaded file {file_path}: {e}")
            raise
    
    def read_word_document(self, file_path: Path) -> str:
        """
        Extract text content from a Word document
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            String containing all text from the document
        """
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    content_parts.append("\n".join(table_data))
            
            content = "\n\n".join(content_parts)
            logger.info(f"Extracted {len(content)} characters from Word document")
            return content
            
        except Exception as e:
            logger.error(f"Error reading Word document: {e}")
            raise
    
    def read_excel_content(self, file_path: Path) -> str:
        """
        Extract content from Excel file (all sheets)
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            String containing all content from Excel sheets
        """
        try:
            xl_file = pd.ExcelFile(file_path)
            content_parts = []
            
            for sheet_name in xl_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                content_parts.append(f"Sheet: {sheet_name}")
                content_parts.append(df.to_string())
            
            content = "\n\n".join(content_parts)
            logger.info(f"Extracted content from {len(xl_file.sheet_names)} Excel sheets")
            return content
            
        except Exception as e:
            logger.error(f"Error reading Excel file: {e}")
            raise
    
    async def clean_csv_with_grok(self, uploaded_df: pd.DataFrame) -> pd.DataFrame:
        """
        Use Grok AI to clean and standardize the uploaded CSV to match the template
        
        Args:
            uploaded_df: DataFrame from uploaded file
            
        Returns:
            Cleaned DataFrame matching the template structure
        """
        try:
            # Get template sample
            template_sample = self.load_template_sample()
            
            # Convert uploaded data to CSV string (limit to first 100 rows for API)
            uploaded_csv = uploaded_df.head(100).to_csv(index=False)
            
            # Format the prompt
            prompt = CSV_CLEANING_PROMPT.format(
                template_name=self.template_name,
                template_columns=", ".join(self.template_columns),
                template_sample=template_sample,
                uploaded_data=uploaded_csv
            )
            
            # Call Grok API
            logger.info("Sending CSV to Grok for cleaning and standardization...")
            response = await asyncio.to_thread(
                self.grok_client.chat.completions.create,
                model=settings.GROK_MODEL,
                messages=[
                    {"role": "system", "content": "You are a data processing expert. Return ONLY valid CSV data without any markdown formatting or explanations."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=8000,
                temperature=0.3  # Low temperature for consistency
            )
            
            cleaned_csv = response.choices[0].message.content.strip()
            
            # Remove markdown code blocks if present
            if cleaned_csv.startswith("```"):
                lines = cleaned_csv.split("\n")
                cleaned_csv = "\n".join([line for line in lines if not line.startswith("```")])
            
            # Parse the cleaned CSV
            cleaned_df = pd.read_csv(StringIO(cleaned_csv))
            
            logger.info(f"Successfully cleaned CSV with Grok. Rows: {len(cleaned_df)}, Columns: {len(cleaned_df.columns)}")
            
            return cleaned_df
            
        except Exception as e:
            logger.error(f"Error cleaning CSV with Grok: {e}")
            # Fallback: return original dataframe
            logger.warning("Falling back to original uploaded data")
            return uploaded_df
    
    async def extract_from_document_with_grok(self, document_content: str) -> pd.DataFrame:
        """
        Use Grok AI to extract data from Word/Excel document content and map to template
        
        Args:
            document_content: Text content from Word/Excel document
            
        Returns:
            DataFrame with extracted data mapped to template structure
        """
        try:
            # Get template sample
            template_sample = self.load_template_sample()
            
            # Format the prompt
            prompt = DOCUMENT_EXTRACTION_PROMPT.format(
                template_name=self.template_name,
                template_columns=", ".join(self.template_columns),
                template_sample=template_sample,
                document_content=document_content[:15000]  # Limit content size
            )
            
            # Call Grok API
            logger.info("Sending document content to Grok for data extraction...")
            response = await asyncio.to_thread(
                self.grok_client.chat.completions.create,
                model=settings.GROK_MODEL,
                messages=[
                    {"role": "system", "content": "You are a data extraction expert. Extract ESG data from documents and return ONLY valid CSV data without any markdown formatting or explanations."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=10000,
                temperature=0.3
            )
            
            extracted_csv = response.choices[0].message.content.strip()
            
            # Remove markdown code blocks if present
            if extracted_csv.startswith("```"):
                lines = extracted_csv.split("\n")
                extracted_csv = "\n".join([line for line in lines if not line.startswith("```")])
            
            # Parse the extracted CSV
            extracted_df = pd.read_csv(StringIO(extracted_csv))
            
            logger.info(f"Successfully extracted data from document with Grok. Rows: {len(extracted_df)}, Columns: {len(extracted_df.columns)}")
            
            return extracted_df
            
        except Exception as e:
            logger.error(f"Error extracting from document with Grok: {e}")
            # Fallback: return empty dataframe with template columns
            logger.warning("Falling back to empty dataframe")
            return pd.DataFrame(columns=self.template_columns)
    
    def create_perfect_match_result(self, cleaned_df: pd.DataFrame) -> ColumnMatchResult:
        """
        Create a ColumnMatchResult indicating perfect match after Grok cleaning
        
        Args:
            cleaned_df: Cleaned DataFrame from Grok
            
        Returns:
            ColumnMatchResult with 100% match
        """
        # Get columns from cleaned dataframe
        cleaned_columns = [col for col in cleaned_df.columns if col and str(col).strip()]
        
        # All template columns should be matched now
        matched = list(set(self.template_columns).intersection(set(cleaned_columns)))
        
        # Extra columns not in template
        extra_columns = list(set(cleaned_columns) - set(self.template_columns))
        
        # Missing columns from template
        missing_columns = list(set(self.template_columns) - set(cleaned_columns))
        
        # Calculate match percentage
        if len(self.template_columns) > 0:
            match_percentage = (len(matched) / len(self.template_columns)) * 100
        else:
            match_percentage = 100.0
        
        has_ambiguity = bool(extra_columns or missing_columns)
        ambiguity_message = None
        
        if has_ambiguity:
            messages = []
            if missing_columns:
                messages.append(f"Missing {len(missing_columns)} template column(s): {', '.join(missing_columns)}")
            if extra_columns:
                messages.append(f"Added {len(extra_columns)} extra column(s): {', '.join(extra_columns)}")
            ambiguity_message = " | ".join(messages)
        
        result = ColumnMatchResult(
            matched_columns=sorted(matched),
            unmatched_uploaded=sorted(extra_columns),
            unmatched_template=sorted(missing_columns),
            match_percentage=round(match_percentage, 2),
            total_uploaded_columns=len(cleaned_columns),
            total_template_columns=len(self.template_columns),
            has_ambiguity=has_ambiguity,
            ambiguity_message=ambiguity_message
        )
        
        logger.info(f"Created match result: {match_percentage}% match, {len(matched)} matched columns")
        return result
    
    def extract_required_data(
        self, 
        cleaned_df: pd.DataFrame,
        match_result: ColumnMatchResult
    ) -> List[Dict[str, Any]]:
        """
        Extract all columns from cleaned data
        
        Args:
            cleaned_df: Cleaned DataFrame from Grok
            match_result: Result from column matching
            
        Returns:
            List of dictionaries containing all extracted data
        """
        # Extract all data from the DataFrame
        extracted_data = []
        
        # Get all column names from the cleaned file
        all_columns = cleaned_df.columns.tolist()
        
        for idx, row in cleaned_df.iterrows():
            record = {}
            # Include all columns from the cleaned file
            for col in all_columns:
                # Skip invalid/unnamed columns
                if col and str(col).strip() and not str(col).startswith('Unnamed:'):
                    # Convert to string and handle NaN values
                    value = row.get(col, '')
                    record[col] = str(value) if pd.notna(value) else ''
            
            extracted_data.append(record)
        
        logger.info(f"Extracted {len(extracted_data)} records with {len(all_columns)} columns from cleaned data")
        return extracted_data
    
    async def extract_columns_only(self, file_path: Path) -> List[str]:
        """
        Extract only column names from uploaded file without full processing
        Supports CSV, Excel (.xlsx, .xls), and Word (.docx) files
        
        Args:
            file_path: Path to uploaded file
            
        Returns:
            List of column names detected in the file
        """
        file_extension = file_path.suffix.lower()
        
        try:
            # Handle Word documents (.docx)
            if file_extension == '.docx':
                logger.info("Extracting columns from Word document...")
                # Read document content
                document_content = self.read_word_document(file_path)
                # Use Grok to identify columns/fields from the document
                columns = await self.extract_columns_from_text(document_content)
            
            # Handle Excel files (.xlsx, .xls)
            elif file_extension in ['.xlsx', '.xls']:
                logger.info("Extracting columns from Excel file...")
                try:
                    # First try to load as structured CSV-like data
                    uploaded_df = self.load_uploaded_file(file_path)
                    # Check if it looks like structured data (has reasonable column names)
                    if len(uploaded_df.columns) > 3 and not all('Unnamed' in str(col) for col in uploaded_df.columns):
                        columns = [col for col in uploaded_df.columns if col and str(col).strip() and not str(col).startswith('Unnamed:')]
                    else:
                        # Treat as document with unstructured content
                        document_content = self.read_excel_content(file_path)
                        columns = await self.extract_columns_from_text(document_content)
                except Exception as e:
                    logger.warning(f"Could not process as structured Excel, treating as document: {e}")
                    document_content = self.read_excel_content(file_path)
                    columns = await self.extract_columns_from_text(document_content)
            
            # Handle CSV files
            elif file_extension == '.csv':
                logger.info("Extracting columns from CSV file...")
                uploaded_df = self.load_uploaded_file(file_path)
                columns = [col for col in uploaded_df.columns if col and str(col).strip() and not str(col).startswith('Unnamed:')]
            
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            logger.info(f"Extracted {len(columns)} columns from file")
            return sorted(columns)
            
        except Exception as e:
            logger.error(f"Error extracting columns: {e}")
            raise
    
    async def extract_columns_from_text(self, text_content: str) -> List[str]:
        """
        Use Grok AI to extract/identify column names/fields from unstructured text
        
        Args:
            text_content: Text content from document
            
        Returns:
            List of identified column/field names
        """
        try:
            prompt = f"""Analyze the following document content and identify all the data fields/columns that appear to contain ESG-related information.

Document Content:
{text_content[:5000]}  # Limit to first 5000 characters

Please extract a list of field names/column headers that represent data points in this document.
Return ONLY a comma-separated list of field names, nothing else.
For example: "Company Name, Year, Total Emissions, Energy Consumption, Water Usage"
"""
            
            logger.info("Using Grok to extract columns from document text...")
            response = await asyncio.to_thread(
                self.grok_client.chat.completions.create,
                model=settings.GROK_MODEL,
                messages=[
                    {"role": "system", "content": "You are a data extraction expert. Extract field names from documents and return them as a simple comma-separated list."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000,
                temperature=0.3
            )
            
            columns_text = response.choices[0].message.content.strip()
            
            # Parse the comma-separated list
            columns = [col.strip() for col in columns_text.split(',') if col.strip()]
            
            logger.info(f"Extracted {len(columns)} columns from text using Grok")
            return columns
            
        except Exception as e:
            logger.error(f"Error extracting columns from text with Grok: {e}")
            return []
    
    async def process_file_with_mappings(
        self,
        file_path: Path,
        mappings: List[Any]
    ) -> Tuple[ColumnMatchResult, List[Dict[str, Any]]]:
        """
        Process file with user-provided column mappings
        
        Args:
            file_path: Path to uploaded file
            mappings: List of ColumnMapping objects with user-confirmed mappings
            
        Returns:
            Tuple of (ColumnMatchResult, extracted_data)
        """
        file_extension = file_path.suffix.lower()
        
        # Create mapping dictionary
        mapping_dict = {}
        for mapping in mappings:
            if mapping.uploaded_column:
                mapping_dict[mapping.template_column] = mapping.uploaded_column
        
        logger.info(f"Processing file with {len(mapping_dict)} user-provided mappings")
        
        # Load the raw file data
        if file_extension == '.docx':
            # For Word documents, extract with Grok then apply mappings
            document_content = self.read_word_document(file_path)
            raw_df = await self.extract_from_document_with_grok(document_content)
        elif file_extension in ['.xlsx', '.xls']:
            try:
                raw_df = self.load_uploaded_file(file_path)
                if len(raw_df.columns) <= 3 or all('Unnamed' in str(col) for col in raw_df.columns):
                    document_content = self.read_excel_content(file_path)
                    raw_df = await self.extract_from_document_with_grok(document_content)
            except Exception:
                document_content = self.read_excel_content(file_path)
                raw_df = await self.extract_from_document_with_grok(document_content)
        elif file_extension == '.csv':
            raw_df = self.load_uploaded_file(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Apply column mappings to create mapped dataframe
        mapped_df = await self.apply_column_mappings(raw_df, mapping_dict)
        
        # Create match result
        match_result = self.create_match_result_from_mappings(mapped_df, mapping_dict)
        
        # Extract data
        extracted_data = self.extract_required_data(mapped_df, match_result)
        
        return match_result, extracted_data
    
    async def apply_column_mappings(
        self,
        raw_df: pd.DataFrame,
        mapping_dict: Dict[str, str]
    ) -> pd.DataFrame:
        """
        Apply user-provided column mappings to raw dataframe
        Uses Grok to intelligently map and transform the data
        
        Args:
            raw_df: Raw DataFrame from uploaded file
            mapping_dict: Dictionary mapping template columns to uploaded columns
            
        Returns:
            DataFrame with columns mapped according to user selection
        """
        try:
            # If the dataframe already has columns similar to template, use direct mapping
            # Otherwise, use Grok to transform the data
            
            # Build the mapped dataframe
            mapped_data = {}
            
            # First, add mapped columns
            for template_col, uploaded_col in mapping_dict.items():
                if uploaded_col in raw_df.columns:
                    mapped_data[template_col] = raw_df[uploaded_col]
                else:
                    # Column not found, create empty column
                    mapped_data[template_col] = pd.Series([None] * len(raw_df))
            
            # Add unmapped template columns as empty
            for template_col in self.template_columns:
                if template_col not in mapped_data:
                    mapped_data[template_col] = pd.Series([None] * len(raw_df))
            
            # Add extra columns from uploaded file that weren't mapped
            for col in raw_df.columns:
                if col not in mapping_dict.values() and col not in mapped_data:
                    mapped_data[col] = raw_df[col]
            
            mapped_df = pd.DataFrame(mapped_data)
            
            logger.info(f"Applied column mappings. Result: {len(mapped_df.columns)} columns, {len(mapped_df)} rows")
            return mapped_df
            
        except Exception as e:
            logger.error(f"Error applying column mappings: {e}")
            raise
    
    def create_match_result_from_mappings(
        self,
        mapped_df: pd.DataFrame,
        mapping_dict: Dict[str, str]
    ) -> ColumnMatchResult:
        """
        Create a ColumnMatchResult based on user mappings
        
        Args:
            mapped_df: DataFrame after applying mappings
            mapping_dict: User-provided column mappings
            
        Returns:
            ColumnMatchResult
        """
        # Get columns from mapped dataframe
        mapped_columns = [col for col in mapped_df.columns if col and str(col).strip()]
        
        # Matched columns are those in the mapping
        matched = list(mapping_dict.keys())
        
        # Extra columns not in template
        extra_columns = list(set(mapped_columns) - set(self.template_columns))
        
        # Missing columns from template (not mapped)
        missing_columns = list(set(self.template_columns) - set(matched))
        
        # Calculate match percentage
        if len(self.template_columns) > 0:
            match_percentage = (len(matched) / len(self.template_columns)) * 100
        else:
            match_percentage = 100.0
        
        has_ambiguity = bool(extra_columns or missing_columns)
        ambiguity_message = None
        
        if has_ambiguity:
            messages = []
            if missing_columns:
                messages.append(f"Missing {len(missing_columns)} template column(s): {', '.join(missing_columns)}")
            if extra_columns:
                messages.append(f"Added {len(extra_columns)} extra column(s): {', '.join(extra_columns)}")
            ambiguity_message = " | ".join(messages)
        
        result = ColumnMatchResult(
            matched_columns=sorted(matched),
            unmatched_uploaded=sorted(extra_columns),
            unmatched_template=sorted(missing_columns),
            match_percentage=round(match_percentage, 2),
            total_uploaded_columns=len(mapped_columns),
            total_template_columns=len(self.template_columns),
            has_ambiguity=has_ambiguity,
            ambiguity_message=ambiguity_message
        )
        
        logger.info(f"Created match result from mappings: {match_percentage}% match, {len(matched)} matched columns")
        return result
    
    async def process_file(self, file_path: Path) -> Tuple[ColumnMatchResult, List[Dict[str, Any]]]:
        """
        Process uploaded file: load, clean with Grok, and extract data
        Supports CSV, Excel (.xlsx, .xls), and Word (.docx) files
        
        Args:
            file_path: Path to uploaded file
            
        Returns:
            Tuple of (ColumnMatchResult, extracted_data)
        """
        file_extension = file_path.suffix.lower()
        
        # Handle Word documents (.docx)
        if file_extension == '.docx':
            logger.info("Processing Word document...")
            # Read document content
            document_content = self.read_word_document(file_path)
            # Extract data using Grok
            cleaned_df = await self.extract_from_document_with_grok(document_content)
        
        # Handle Excel files (.xlsx, .xls) - try as structured data first, then as document
        elif file_extension in ['.xlsx', '.xls']:
            logger.info("Processing Excel file...")
            try:
                # First try to load as structured CSV-like data
                uploaded_df = self.load_uploaded_file(file_path)
                # Check if it looks like structured data (has reasonable column names)
                if len(uploaded_df.columns) > 3 and not all('Unnamed' in str(col) for col in uploaded_df.columns):
                    cleaned_df = await self.clean_csv_with_grok(uploaded_df)
                else:
                    # Treat as document with unstructured content
                    document_content = self.read_excel_content(file_path)
                    cleaned_df = await self.extract_from_document_with_grok(document_content)
            except Exception as e:
                logger.warning(f"Could not process as structured Excel, treating as document: {e}")
                document_content = self.read_excel_content(file_path)
                cleaned_df = await self.extract_from_document_with_grok(document_content)
        
        # Handle CSV files
        elif file_extension == '.csv':
            logger.info("Processing CSV file...")
            uploaded_df = self.load_uploaded_file(file_path)
            cleaned_df = await self.clean_csv_with_grok(uploaded_df)
        
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Create match result
        match_result = self.create_perfect_match_result(cleaned_df)
        
        # Extract data
        extracted_data = self.extract_required_data(cleaned_df, match_result)
        
        return match_result, extracted_data


def get_data_summary(extracted_data: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Generate a summary of extracted data for reporting
    
    Args:
        extracted_data: List of extracted data records
        
    Returns:
        Dictionary containing data summary
    """
    if not extracted_data:
        return {}
    
    # Define possible column name mappings
    SECTION_COLUMNS = ["Section / القسم", "Section (EN)", "Section (AR)", "section"]
    CURRENT_COLUMNS = ["Current", "Current / العام الحالي", "Response / الإدخال", "current"]
    TARGET_COLUMNS = ["Target", "Target / الهدف", "target"]
    
    # Extract sections
    sections = []
    section_counts = {}
    filled_current = 0
    filled_target = 0
    
    for row in extracted_data:
        # Find section
        section = _find_column(row, SECTION_COLUMNS)
        if section:
            sections.append(section)
            section_counts[section] = section_counts.get(section, 0) + 1
        
        # Check filled fields
        if _find_column(row, CURRENT_COLUMNS):
            filled_current += 1
        if _find_column(row, TARGET_COLUMNS):
            filled_target += 1
    
    # Get unique sections
    unique_sections = list(set(sections))
    total_fields = len(extracted_data)
    
    summary = {
        'total_records': total_fields,
        'sections': unique_sections,
        'section_counts': section_counts,
        'total_fields': total_fields,
        'filled_current': filled_current,
        'filled_target': filled_target,
        'completion_rate_current': round((filled_current / total_fields) * 100, 2) if total_fields > 0 else 0,
        'completion_rate_target': round((filled_target / total_fields) * 100, 2) if total_fields > 0 else 0
    }
    
    return summary


def _find_column(row: Dict[str, Any], possible_names: List[str]) -> Any:
    """
    Find a column value by checking multiple possible column names
    
    Args:
        row: Data row dictionary
        possible_names: List of possible column names to check
        
    Returns:
        Value if found, None otherwise
    """
    for name in possible_names:
        if name in row:
            value = row[name]
            if pd.notna(value) and str(value).strip() and str(value) != 'nan':
                return value
    return None


def format_data_for_report(extracted_data: List[Dict[str, Any]]) -> str:
    """
    Format extracted data into a readable string for AI report generation
    
    Args:
        extracted_data: List of extracted data records
        
    Returns:
        Formatted string representation of the data
    """
    if not extracted_data:
        return "No data available."
    
    formatted_lines = []
    formatted_lines.append("ESG DATA SUMMARY")
    formatted_lines.append("=" * 80)
    formatted_lines.append("")
    
    # Define possible column name mappings for different templates
    SECTION_COLUMNS = ["Section / القسم", "Section (EN)", "Section (AR)", "section"]
    FIELD_COLUMNS = ["Field (EN)", "الحقل (AR)", "field"]
    PREV_YEAR_COLUMNS = ["Prev Year", "Prev Year / العام السابق", "prev_year"]
    CURRENT_COLUMNS = ["Current", "Current / العام الحالي", "Response / الإدخال", "current"]
    TARGET_COLUMNS = ["Target", "Target / الهدف", "target"]
    UNIT_COLUMNS = ["Unit", "Unit / الوحدة", "unit"]
    NOTES_COLUMNS = ["Notes", "Notes / ملاحظات", "notes"]
    
    # Group data by section
    current_section = None
    
    for row in extracted_data:
        # Find section value
        section = _find_column(row, SECTION_COLUMNS)
        
        if section and section != current_section:
            current_section = section
            formatted_lines.append(f"\n## {section}")
            formatted_lines.append("-" * 80)
        
        # Find field value
        field = _find_column(row, FIELD_COLUMNS)
        
        if field:
            formatted_lines.append(f"\n### {field}")
            
            # Find other values
            prev_year = _find_column(row, PREV_YEAR_COLUMNS)
            if prev_year:
                formatted_lines.append(f"  Previous Year: {prev_year}")
            
            current = _find_column(row, CURRENT_COLUMNS)
            if current:
                formatted_lines.append(f"  Current: {current}")
            
            target = _find_column(row, TARGET_COLUMNS)
            if target:
                formatted_lines.append(f"  Target: {target}")
            
            unit = _find_column(row, UNIT_COLUMNS)
            if unit:
                formatted_lines.append(f"  Unit: {unit}")
            
            notes = _find_column(row, NOTES_COLUMNS)
            if notes:
                formatted_lines.append(f"  Notes: {notes}")
    
    return "\n".join(formatted_lines)


def calculate_change_analysis(extracted_data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Calculate change percentage and status for each record
    
    Args:
        extracted_data: List of extracted data records
        
    Returns:
        List of records with change analysis added
    """
    # Define possible column name mappings
    PREV_YEAR_COLUMNS = ["Prev Year", "Prev Year / العام السابق", "prev_year"]
    CURRENT_COLUMNS = ["Current", "Current / العام الحالي", "Response / الإدخال", "current"]
    FIELD_COLUMNS = ["Field (EN)", "الحقل (AR)", "field"]
    
    # Keywords indicating "lower is better" metrics
    LOWER_IS_BETTER = [
        'emission', 'ghg', 'co2', 'waste', 'discharge', 'consumption',
        'intensity', 'turnover', 'accident', 'incident', 'gap', 'pay gap'
    ]
    
    analyzed_data = []
    
    for row in extracted_data:
        analysis = {
            "field_data": row,
            "change_percentage": None,
            "change_status": None
        }
        
        # Get prev year and current values
        prev_year_value = _find_column(row, PREV_YEAR_COLUMNS)
        current_value = _find_column(row, CURRENT_COLUMNS)
        field_name = _find_column(row, FIELD_COLUMNS) or ""
        
        if prev_year_value and current_value:
            try:
                # Try to convert to float
                prev = float(str(prev_year_value).replace(',', '').replace('%', '').strip())
                curr = float(str(current_value).replace(',', '').replace('%', '').strip())
                
                # Calculate percentage change
                if prev != 0:
                    change = ((curr - prev) / abs(prev)) * 100
                    analysis["change_percentage"] = round(change, 2)
                    
                    # Determine if lower is better for this field
                    lower_is_better = any(keyword in field_name.lower() for keyword in LOWER_IS_BETTER)
                    
                    # Determine status based on change and field type
                    if abs(change) == 0:  # Within 5% considered slight
                        analysis["change_status"] = "unchanged"
                    if abs(change) <= 5:  # Within 5% considered slight
                        analysis["change_status"] = "slight"
                    elif change > 5:  # Increased
                        analysis["change_status"] = "worsened" if lower_is_better else "improved"
                    else:  # Decreased (change < -5)
                        analysis["change_status"] = "improved" if lower_is_better else "worsened"
                else:
                    # Previous year was 0
                    if curr > 0:
                        analysis["change_percentage"] = 100.0
                        lower_is_better = any(keyword in field_name.lower() for keyword in LOWER_IS_BETTER)
                        analysis["change_status"] = "worsened" if lower_is_better else "improved"
                    elif curr == 0:
                        analysis["change_percentage"] = 0.0
                        analysis["change_status"] = "slight"
                    
            except (ValueError, TypeError):
                # Can't convert to numbers, skip analysis
                pass
        
        analyzed_data.append(analysis)
    
    return analyzed_data
