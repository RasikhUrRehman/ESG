"""
Report generation using Grok AI and document formatting
"""
import os
import json
import logging
import asyncio
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime
from openai import OpenAI

# Document generation libraries
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd

from .config import settings
from .prompts import get_report_prompt, SYSTEM_PROMPT, CHART_INSTRUCTIONS

logger = logging.getLogger(__name__)


class GrokAPIClient:
    """Client for interacting with Grok AI API using OpenAI-compatible endpoint"""
    
    def __init__(self):
        self.api_key = settings.GROK_API_KEY
        self.base_url = settings.GROK_API_BASE
        self.model = settings.GROK_MODEL
        
        # Initialize OpenAI client with xAI Grok endpoint
        self.client = OpenAI(
            api_key=self.api_key,
            base_url=self.base_url,
        )
        
    async def generate_content(
        self, 
        prompt: str, 
        system_prompt: str = SYSTEM_PROMPT,
        max_tokens: int = 4000,
        temperature: float = 0.7
    ) -> str:
        """
        Generate content using Grok AI
        
        Args:
            prompt: User prompt
            system_prompt: System prompt for context
            max_tokens: Maximum tokens in response
            temperature: Controls randomness (0.0-2.0)
            
        Returns:
            Generated content as string
            
        Raises:
            Exception: If API request fails
        """
        try:
            # Run the synchronous OpenAI client in a thread pool to avoid blocking
            response = await asyncio.to_thread(
                self.client.chat.completions.create,
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_tokens,
                temperature=temperature
            )
            
            content = response.choices[0].message.content
            
            logger.info("Successfully generated content using Grok AI")
            return content
            
        except Exception as e:
            logger.error(f"Error generating content with Grok AI: {e}")
            
            # Enhanced error handling
            error_msg = str(e)
            if "401" in error_msg:
                logger.error("Invalid or expired API key")
                raise Exception("Invalid or expired Grok API key")
            elif "429" in error_msg:
                logger.error("Rate limit exceeded")
                raise Exception("Rate limit exceeded. Please try again later.")
            elif "404" in error_msg and "model" in error_msg.lower():
                logger.error(f"Model '{self.model}' not found")
                raise Exception(f"Model '{self.model}' not found. Check available models.")
            else:
                raise


class ChartGenerator:
    """Generate charts for ESG reports"""
    
    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(exist_ok=True)
        sns.set_style("whitegrid")
    
    def _extract_numeric_data(self, data: List[Dict[str, Any]], keywords: List[str]) -> Dict[str, float]:
        """Extract numeric data for fields matching keywords"""
        result = {}
        for row in data:
            for key, value in row.items():
                field_name = row.get('Field (EN)', row.get('field', str(key)))
                
                # Check if field matches keywords
                if any(term in str(field_name).lower() for term in keywords):
                    current_value = row.get('Current', row.get('current', row.get('Response / الإدخال', '')))
                    
                    try:
                        if current_value and str(current_value).strip() and str(current_value).lower() not in ['nan', 'not available', 'n/a', '']:
                            numeric_value = float(str(current_value).replace(',', '').replace('%', ''))
                            if numeric_value > 0:
                                result[field_name] = numeric_value
                    except (ValueError, TypeError):
                        continue
        return result
    
    def analyze_data_for_charts(self, data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Analyze data to determine what charts can be created
        Returns list of chart specifications
        """
        chart_specs = []
        
        # Check for emissions data
        emissions = self._extract_numeric_data(data, ['scope', 'emission', 'ghg', 'co2'])
        if len(emissions) >= 2:
            chart_specs.append({
                'type': 'bar',
                'category': 'emissions',
                'title': 'Greenhouse Gas Emissions',
                'data': emissions,
                'ylabel': 'Emissions (tCO2e)'
            })
        
        # Check for energy data
        energy = self._extract_numeric_data(data, ['energy', 'renewable', 'electricity', 'fuel'])
        if len(energy) >= 2:
            chart_specs.append({
                'type': 'bar',
                'category': 'energy',
                'title': 'Energy Consumption',
                'data': energy,
                'ylabel': 'Energy (units vary)'
            })
        
        # Check for diversity/gender data
        diversity = self._extract_numeric_data(data, ['diversity', 'gender', 'female', 'male', 'women', 'board'])
        if len(diversity) >= 2:
            chart_specs.append({
                'type': 'pie',
                'category': 'diversity',
                'title': 'Diversity Metrics',
                'data': diversity
            })
        
        # Check for social metrics
        social = self._extract_numeric_data(data, ['employee', 'turnover', 'safety', 'injury', 'training'])
        if len(social) >= 2:
            chart_specs.append({
                'type': 'bar',
                'category': 'social',
                'title': 'Social Metrics',
                'data': social,
                'ylabel': 'Value'
            })
        
        # Check for water data
        water = self._extract_numeric_data(data, ['water', 'consumption', 'reclamation', 'discharge'])
        if len(water) >= 2:
            chart_specs.append({
                'type': 'bar',
                'category': 'water',
                'title': 'Water Management',
                'data': water,
                'ylabel': 'Volume (units vary)'
            })
        
        # Check for waste data
        waste = self._extract_numeric_data(data, ['waste', 'recycling', 'landfill'])
        if len(waste) >= 2:
            chart_specs.append({
                'type': 'bar',
                'category': 'waste',
                'title': 'Waste Management',
                'data': waste,
                'ylabel': 'Amount (units vary)'
            })
        
        # Check for trend data (requires prev year, current, target)
        trend_data = self._check_for_trends(data)
        for trend_spec in trend_data:
            chart_specs.append(trend_spec)
        
        return chart_specs
    
    def _check_for_trends(self, data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Check for metrics with historical/trend data"""
        trends = []
        
        for row in data:
            field_name = row.get('Field (EN)', row.get('field', ''))
            if not field_name:
                continue
            
            prev_year = row.get('Prev Year', row.get('prev_year', row.get('Prev Year / العام السابق', '')))
            current = row.get('Current', row.get('current', row.get('Current / العام الحالي', '')))
            target = row.get('Target', row.get('target', row.get('Target / الهدف', '')))
            
            values = []
            periods = []
            
            try:
                if prev_year and str(prev_year).strip() and str(prev_year).lower() not in ['nan', 'not available', 'n/a', '']:
                    val = float(str(prev_year).replace(',', ''))
                    if val > 0:
                        values.append(val)
                        periods.append('Previous Year')
                
                if current and str(current).strip() and str(current).lower() not in ['nan', 'not available', 'n/a', '']:
                    val = float(str(current).replace(',', ''))
                    if val > 0:
                        values.append(val)
                        periods.append('Current')
                
                if target and str(target).strip() and str(target).lower() not in ['nan', 'not available', 'n/a', '']:
                    val = float(str(target).replace(',', ''))
                    if val > 0:
                        values.append(val)
                        periods.append('Target')
                
                if len(values) >= 2:
                    trends.append({
                        'type': 'line',
                        'category': 'trend',
                        'title': f'{field_name} - Trend',
                        'data': {period: val for period, val in zip(periods, values)},
                        'ylabel': 'Value',
                        'field_name': field_name
                    })
            except (ValueError, TypeError):
                continue
        
        return trends[:3]  # Limit to 3 trend charts to avoid cluttering
    
    def create_chart(self, chart_spec: Dict[str, Any], filename: str) -> Path:
        """Create a chart based on specification"""
        chart_type = chart_spec['type']
        
        if chart_type == 'bar':
            return self._create_bar_chart(chart_spec, filename)
        elif chart_type == 'pie':
            return self._create_pie_chart(chart_spec, filename)
        elif chart_type == 'line':
            return self._create_line_chart(chart_spec, filename)
        else:
            logger.warning(f"Unknown chart type: {chart_type}")
            return None
    
    def _create_bar_chart(self, spec: Dict[str, Any], filename: str) -> Path:
        """Create bar chart"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        data = spec['data']
        labels = list(data.keys())
        values = list(data.values())
        
        colors_list = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f']
        bar_colors = colors_list[:len(labels)]
        
        ax.bar(labels, values, color=bar_colors)
        ax.set_ylabel(spec.get('ylabel', 'Value'), fontsize=12)
        ax.set_title(spec['title'], fontsize=14, fontweight='bold')
        plt.xticks(rotation=45, ha='right')
        ax.grid(axis='y', alpha=0.3)
        
        plt.tight_layout()
        
        chart_path = self.output_dir / filename
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        logger.info(f"Created bar chart: {spec['title']} with {len(data)} data points")
        return chart_path
    
    def _create_pie_chart(self, spec: Dict[str, Any], filename: str) -> Path:
        """Create pie chart"""
        fig, ax = plt.subplots(figsize=(8, 8))
        
        data = spec['data']
        labels = list(data.keys())
        sizes = list(data.values())
        
        colors_list = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99', '#ff99cc', '#c2c2f0']
        pie_colors = colors_list[:len(labels)]
        
        ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=pie_colors)
        ax.set_title(spec['title'], fontsize=14, fontweight='bold')
        
        chart_path = self.output_dir / filename
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        logger.info(f"Created pie chart: {spec['title']} with {len(data)} data points")
        return chart_path
    
    def _create_line_chart(self, spec: Dict[str, Any], filename: str) -> Path:
        """Create line chart for trends"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        data = spec['data']
        periods = list(data.keys())
        values = list(data.values())
        
        ax.plot(periods, values, marker='o', linewidth=2, markersize=8, color='#1f77b4')
        ax.set_ylabel(spec.get('ylabel', 'Value'), fontsize=12)
        ax.set_title(spec['title'], fontsize=14, fontweight='bold')
        ax.grid(True, alpha=0.3)
        plt.xticks(rotation=0)
        
        plt.tight_layout()
        
        chart_path = self.output_dir / filename
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        logger.info(f"Created line chart: {spec['title']} with {len(data)} data points")
        return chart_path
        
    def create_emissions_chart(self, data: List[Dict[str, Any]], filename: str) -> Path:
        """Create emissions comparison chart"""
        # Extract emissions data from the ESG data
        emissions = {}
        for row in data:
            # Look for emission-related fields
            for key, value in row.items():
                if any(term in str(key).lower() for term in ['scope', 'emission', 'ghg', 'co2']):
                    field_name = row.get('Field (EN)', row.get('field', key))
                    current_value = row.get('Current', row.get('current', row.get('Response / الإدخال', '')))
                    
                    # Try to convert to numeric
                    try:
                        if current_value and str(current_value).strip() and str(current_value) != 'nan':
                            numeric_value = float(str(current_value).replace(',', ''))
                            if numeric_value > 0:  # Only include positive values
                                emissions[field_name] = numeric_value
                    except (ValueError, TypeError):
                        continue
        
        # Don't create chart if no data
        if not emissions:
            logger.info("No emissions data available - skipping emissions chart")
            return None
        
        fig, ax = plt.subplots(figsize=(10, 6))
        labels = list(emissions.keys())
        values = list(emissions.values())
        
        colors_list = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b']
        bar_colors = colors_list[:len(labels)]
        
        ax.bar(labels, values, color=bar_colors)
        ax.set_ylabel('Emissions (tCO2e)', fontsize=12)
        ax.set_title('Greenhouse Gas Emissions', fontsize=14, fontweight='bold')
        plt.xticks(rotation=45, ha='right')
        ax.grid(axis='y', alpha=0.3)
        
        plt.tight_layout()
        
        chart_path = self.output_dir / filename
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        logger.info(f"Created emissions chart with {len(emissions)} data points")
        return chart_path
    
    def create_diversity_chart(self, data: List[Dict[str, Any]], filename: str) -> Path:
        """Create diversity metrics chart"""
        # Extract diversity data
        diversity_data = {}
        for row in data:
            for key, value in row.items():
                if any(term in str(key).lower() for term in ['diversity', 'gender', 'female', 'male', 'employee']):
                    field_name = row.get('Field (EN)', row.get('field', key))
                    current_value = row.get('Current', row.get('current', row.get('Response / الإدخال', '')))
                    
                    try:
                        if current_value and str(current_value).strip() and str(current_value) != 'nan':
                            numeric_value = float(str(current_value).replace(',', '').replace('%', ''))
                            if numeric_value > 0:  # Only include positive values
                                diversity_data[field_name] = numeric_value
                    except (ValueError, TypeError):
                        continue
        
        # Don't create chart if no data
        if not diversity_data or len(diversity_data) == 0:
            logger.info("No diversity data available - skipping diversity chart")
            return None
        
        fig, ax = plt.subplots(figsize=(8, 8))
        labels = list(diversity_data.keys())
        sizes = list(diversity_data.values())
        
        colors_list = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99', '#ff99cc']
        pie_colors = colors_list[:len(labels)]
        
        ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=pie_colors)
        ax.set_title('Diversity Metrics', fontsize=14, fontweight='bold')
        
        chart_path = self.output_dir / filename
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        logger.info(f"Created diversity chart with {len(diversity_data)} data points")
        return chart_path
    
    def create_trend_chart(self, data: List[Dict[str, Any]], metric: str, filename: str) -> Path:
        """Create trend line chart"""
        # Extract trend data - previous year, current, target
        trends = []
        labels = []
        
        for row in data:
            field_name = row.get('Field (EN)', row.get('field', ''))
            
            if metric.lower() in str(field_name).lower():
                prev_year = row.get('Prev Year', row.get('prev_year', row.get('Prev Year / العام السابق', '')))
                current = row.get('Current', row.get('current', row.get('Current / العام الحالي', '')))
                target = row.get('Target', row.get('target', row.get('Target / الهدف', '')))
                
                try:
                    values = []
                    periods = []
                    
                    if prev_year and str(prev_year).strip() and str(prev_year) != 'nan':
                        val = float(str(prev_year).replace(',', ''))
                        if val > 0:
                            values.append(val)
                            periods.append('Previous Year')
                    
                    if current and str(current).strip() and str(current) != 'nan':
                        val = float(str(current).replace(',', ''))
                        if val > 0:
                            values.append(val)
                            periods.append('Current')
                    
                    if target and str(target).strip() and str(target) != 'nan':
                        val = float(str(target).replace(',', ''))
                        if val > 0:
                            values.append(val)
                            periods.append('Target')
                    
                    if values and len(values) >= 2:  # Need at least 2 points for a trend
                        trends.append((field_name, periods, values))
                except (ValueError, TypeError):
                    continue
        
        # Don't create chart if no trend data
        if not trends:
            logger.info(f"No trend data available for {metric} - skipping trend chart")
            return None
        
        fig, ax = plt.subplots(figsize=(10, 6))
        
        for field_name, periods, values in trends:
            ax.plot(periods, values, marker='o', linewidth=2, label=field_name)
        
        ax.set_ylabel('Value', fontsize=12)
        ax.set_title(f'{metric} Trend Analysis', fontsize=14, fontweight='bold')
        ax.grid(True, alpha=0.3)
        ax.legend(loc='best', fontsize=9)
        plt.xticks(rotation=0)
        plt.tight_layout()
        
        chart_path = self.output_dir / filename
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        logger.info(f"Created trend chart for {metric} with {len(trends)} trends")
        return chart_path
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return chart_path


class PDFReportGenerator:
    """Generate PDF reports using ReportLab"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
        
    def _setup_custom_styles(self):
        """Setup custom paragraph styles"""
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=30,
            alignment=TA_CENTER
        ))
        
        self.styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2ca02c'),
            spaceAfter=12,
            spaceBefore=12
        ))
    
    def _convert_markdown_to_html(self, text: str) -> str:
        """Convert markdown formatting to HTML for ReportLab"""
        import re
        # Convert **text** to <b>text</b>
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        # Convert *text* to <i>text</i>
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        return text
    
    def generate(
        self, 
        content: str, 
        output_path: Path,
        title: str = "ESG Report",
        charts: List[Path] = None
    ) -> Path:
        """
        Generate PDF report
        
        Args:
            content: Report content (markdown or plain text)
            output_path: Path to save PDF
            title: Report title
            charts: List of chart image paths to include
            
        Returns:
            Path to generated PDF
        """
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        story = []
        
        # Title page
        title_para = Paragraph(title, self.styles['CustomTitle'])
        story.append(title_para)
        story.append(Spacer(1, 0.3 * inch))
        
        date_para = Paragraph(
            f"Generated on: {datetime.now().strftime('%B %d, %Y')}",
            self.styles['Normal']
        )
        story.append(date_para)
        story.append(PageBreak())
        
        # Process content
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                story.append(Spacer(1, 0.1 * inch))
                i += 1
                continue
            
            # Check if this is a markdown table
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                # Parse markdown table
                table_lines = [line]
                i += 1
                
                # Get all table rows
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1
                
                # Convert to ReportLab table
                table_data = []
                for row in table_lines:
                    if '---' in row or '|-' in row:  # Skip separator line
                        continue
                    cells = [cell.strip() for cell in row.split('|')]
                    cells = [c for c in cells if c]  # Remove empty cells
                    if cells:
                        table_data.append(cells)
                
                if table_data:
                    # Create table
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2ca02c')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 10),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 0.2 * inch))
                continue
            
            # Handle markdown headers
            if line.startswith('# '):
                para = Paragraph(line[2:], self.styles['CustomTitle'])
                story.append(para)
            elif line.startswith('## '):
                para = Paragraph(line[3:], self.styles['CustomHeading'])
                story.append(para)
            elif line.startswith('### '):
                para = Paragraph(line[4:], self.styles['Heading3'])
                story.append(para)
            elif line.startswith('- ') or line.startswith('* '):
                # Handle list items with inline bold
                text = line[2:]
                text = self._convert_markdown_to_html(text)
                para = Paragraph(text, self.styles['Normal'])
                story.append(para)
            else:
                # Convert markdown bold to HTML
                text = self._convert_markdown_to_html(line)
                para = Paragraph(text, self.styles['Normal'])
                story.append(para)
            
            i += 1
        
        # Add charts if provided
        if charts:
            story.append(PageBreak())
            story.append(Paragraph("Charts and Visualizations", self.styles['CustomHeading']))
            
            for chart_path in charts:
                if chart_path.exists():
                    try:
                        img = Image(str(chart_path), width=6*inch, height=4*inch)
                        story.append(img)
                        story.append(Spacer(1, 0.2 * inch))
                    except Exception as e:
                        logger.error(f"Error adding chart {chart_path}: {e}")
        
        # Build PDF
        doc.build(story)
        logger.info(f"PDF report generated: {output_path}")
        
        return output_path


class WordReportGenerator:
    """Generate Word documents using python-docx"""
    
    def _add_paragraph_with_formatting(self, doc, text):
        """Add paragraph with inline markdown formatting (bold, italic)"""
        import re
        
        # Create new paragraph
        p = doc.add_paragraph()
        
        # Split text by markdown patterns
        # Pattern to match **bold** and *italic*
        pattern = r'(\*\*.*?\*\*|\*.*?\*)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
                
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                p.add_run(part[2:-2]).bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                p.add_run(part[1:-1]).italic = True
            else:
                # Normal text
                p.add_run(part)
        
        return p
    
    def generate(
        self,
        content: str,
        output_path: Path,
        title: str = "ESG Report",
        charts: List[Path] = None
    ) -> Path:
        """
        Generate Word report
        
        Args:
            content: Report content (markdown or plain text)
            output_path: Path to save DOCX
            title: Report title
            charts: List of chart image paths to include
            
        Returns:
            Path to generated DOCX
        """
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Process content
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                doc.add_paragraph()
                i += 1
                continue
            
            # Check if this is a markdown table
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                # Parse markdown table
                table_lines = [line]
                i += 1
                
                # Get all table rows
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1
                
                # Convert to Word table
                table_data = []
                for row in table_lines:
                    if '---' in row or '|-' in row:  # Skip separator line
                        continue
                    cells = [cell.strip() for cell in row.split('|')]
                    cells = [c for c in cells if c]  # Remove empty cells
                    if cells:
                        table_data.append(cells)
                
                if table_data:
                    # Create table
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    # Fill table
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text
                            
                            # Make header row bold
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                    
                    doc.add_paragraph()  # Add space after table
                continue
            
            # Handle markdown headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                # List items with inline formatting
                text = line[2:]
                self._add_paragraph_with_formatting(doc, text).style = 'List Bullet'
            else:
                # Regular paragraph with inline formatting
                self._add_paragraph_with_formatting(doc, line)
            
            i += 1
        
        # Add charts if provided
        if charts:
            doc.add_page_break()
            doc.add_heading("Charts and Visualizations", level=1)
            
            for chart_path in charts:
                if chart_path.exists():
                    try:
                        doc.add_picture(str(chart_path), width=Inches(6))
                        doc.add_paragraph()
                    except Exception as e:
                        logger.error(f"Error adding chart {chart_path}: {e}")
        
        # Save document
        doc.save(str(output_path))
        logger.info(f"Word report generated: {output_path}")
        
        return output_path


class ReportGenerator:
    """Main report generator orchestrator"""
    
    def __init__(self):
        self.grok_client = GrokAPIClient()
        self.pdf_generator = PDFReportGenerator()
        self.word_generator = WordReportGenerator()
        
    async def generate_report(
        self,
        data: List[Dict[str, Any]],
        report_type: str,
        output_format: str,
        output_filename: str,
        include_charts: bool = True
    ) -> Path:
        """
        Generate complete ESG report
        
        Args:
            data: Extracted ESG data
            report_type: Type of report (comprehensive, environmental, etc.)
            output_format: Format (pdf or docx)
            output_filename: Output filename without extension
            include_charts: Whether to include charts
            
        Returns:
            Path to generated report
        """
        from .column_matcher import format_data_for_report
        
        # Format data for AI
        formatted_data = format_data_for_report(data)
        
        # Get appropriate prompt
        prompt = get_report_prompt(report_type, formatted_data)
        
        # Do NOT add CHART_INSTRUCTIONS - we don't want chart suggestions in the text
        # Charts will be generated and added separately
        
        # Generate content using Grok AI
        logger.info(f"Generating {report_type} report content using Grok AI...")
        content = await self.grok_client.generate_content(prompt)
        
        # Generate charts if requested
        charts = []
        if include_charts:
            chart_dir = settings.REPORTS_DIR / "charts"
            chart_generator = ChartGenerator(chart_dir)
            
            try:
                logger.info("Analyzing data to determine which charts to create...")
                
                # Analyze data and get chart specifications
                chart_specs = chart_generator.analyze_data_for_charts(data)
                
                if chart_specs:
                    logger.info(f"Found {len(chart_specs)} charts that can be created")
                    
                    # Create each chart based on specification
                    for i, spec in enumerate(chart_specs):
                        chart_filename = f"{spec['category']}_{i}_{output_filename}.png"
                        chart_path = chart_generator.create_chart(spec, chart_filename)
                        
                        if chart_path and chart_path.exists():
                            charts.append(chart_path)
                    
                    logger.info(f"Successfully created {len(charts)} charts")
                else:
                    logger.info("No sufficient data available for chart generation")
                    
            except Exception as e:
                logger.error(f"Error generating charts: {e}")
        
        # Generate document
        output_path = settings.REPORTS_DIR / f"{output_filename}.{output_format}"
        
        if output_format == "pdf":
            result_path = self.pdf_generator.generate(
                content,
                output_path,
                title=f"{report_type.title()} ESG Report",
                charts=charts
            )
        elif output_format == "docx":
            result_path = self.word_generator.generate(
                content,
                output_path,
                title=f"{report_type.title()} ESG Report",
                charts=charts
            )
        else:
            raise ValueError(f"Unsupported format: {output_format}")
        
        return result_path
